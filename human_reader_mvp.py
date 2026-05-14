"""
Human Reader MVP
AAEL Labs prototype

Goal:
Create a simple, non-coder-friendly screen/document reader that:
1. Reads webpages, PDFs, Word documents, and plain text files.
2. Attempts to clean out citations, footnotes, captions, headers, URLs, and other listening junk.
3. Lets the user start reading from a chosen phrase.
4. Generates browser-playable audio with pause/stop/seek controls.
5. Keeps local pyttsx3 as a backup option, but browser audio is the safer default.

Install requirements:
    pip install streamlit beautifulsoup4 requests pymupdf python-docx readability-lxml edge-tts pyttsx3

Run:
    python -m streamlit run human_reader_mvp.py

Why audio generation instead of only pyttsx3 live reading?
- Streamlit reruns the script when buttons are clicked.
- pyttsx3 uses a blocking speech loop and can throw: RuntimeError: run loop already started.
- Browser audio gives the user normal controls: play, pause, stop, restart, and seek.

Notes:
- This is an MVP, not a full accessibility replacement.
- PDF column detection is basic. It sorts text blocks top-to-bottom and left-to-right.
- For scanned PDFs, OCR would need to be added later.
"""

import asyncio
import os
import subprocess
import sys
import webbrowser
import re
import tempfile
from pathlib import Path
from typing import List, Optional

import requests
import streamlit as st
import fitz  # PyMuPDF
from bs4 import BeautifulSoup
from docx import Document

try:
    import edge_tts
    EDGE_TTS_AVAILABLE = True
except ImportError:
    EDGE_TTS_AVAILABLE = False

try:
    import pyttsx3
    PYTTSX3_AVAILABLE = True
except ImportError:
    PYTTSX3_AVAILABLE = False

try:
    from readability import Document as ReadabilityDocument
    READABILITY_AVAILABLE = True
except ImportError:
    READABILITY_AVAILABLE = False


# -----------------------------
# Text extraction functions
# -----------------------------

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a PDF while attempting to preserve reading order."""
    doc = fitz.open(file_path)
    pages_text = []

    for page in doc:
        blocks = page.get_text("blocks")
        blocks = sorted(blocks, key=lambda b: (round(b[1], 1), round(b[0], 1)))

        page_text = []
        for block in blocks:
            text = block[4].strip()
            if text:
                page_text.append(text)

        pages_text.append("\n".join(page_text))

    return "\n\n".join(pages_text)


def extract_text_from_docx(file_path: str) -> str:
    """Extract paragraph text from a Word document."""
    doc = Document(file_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(paragraphs)


def extract_text_from_txt(file_path: str) -> str:
    """Read text from a plain text file."""
    return Path(file_path).read_text(encoding="utf-8", errors="ignore")


def extract_text_from_webpage(url: str) -> str:
    """Extract main article text from a webpage."""
    headers = {"User-Agent": "Mozilla/5.0 HumanReaderMVP/1.0"}
    response = requests.get(url, headers=headers, timeout=20)
    response.raise_for_status()

    html = response.text

    if READABILITY_AVAILABLE:
        readable = ReadabilityDocument(html)
        html = readable.summary()

    soup = BeautifulSoup(html, "html.parser")

    for tag in soup(["script", "style", "nav", "footer", "header", "aside", "form", "button"]):
        tag.decompose()

    for tag in soup.find_all(["figcaption", "caption"]):
        tag.decompose()

    text_parts = []
    for element in soup.find_all(["h1", "h2", "h3", "p", "li"]):
        text = element.get_text(" ", strip=True)
        if text:
            text_parts.append(text)

    return "\n\n".join(text_parts)


# -----------------------------
# Cleaning functions
# -----------------------------

def remove_urls(text: str) -> str:
    return re.sub(r"https?://\S+|www\.\S+", "", text)


def remove_inline_citations(text: str) -> str:
    patterns = [
        r"\((?:[A-Z][A-Za-z\-]+(?: et al\.)?,?\s*)?\d{4}[a-z]?\)",
        r"\([A-Z][A-Za-z\-]+(?: & [A-Z][A-Za-z\-]+)?, \d{4}[a-z]?\)",
        r"\[[0-9,\s\-]+\]",
        r"\(p\.\s?\d+\)",
        r"\(pp\.\s?\d+\s?-\s?\d+\)",
    ]
    for pattern in patterns:
        text = re.sub(pattern, "", text)
    return text


def remove_footnote_markers(text: str) -> str:
    text = re.sub(r"(?<=\w)\d{1,2}(?=\s)", "", text)
    text = re.sub(r"[\*†‡]+", "", text)
    return text


def remove_captions(text: str) -> str:
    lines = text.splitlines()
    cleaned_lines = []

    caption_starts = (
        "figure ", "fig. ", "table ", "image ", "photo ",
        "source:", "caption:", "credit:", "note:"
    )

    for line in lines:
        stripped = line.strip()
        lower = stripped.lower()

        if any(lower.startswith(start) for start in caption_starts):
            continue

        cleaned_lines.append(line)

    return "\n".join(cleaned_lines)


def remove_metadata_lines(
    text: str,
    remove_titles: bool,
    remove_authors: bool,
    remove_dates: bool,
    remove_footnotes: bool,
) -> str:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    cleaned = []

    date_pattern = re.compile(
        r"\b(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|"
        r"Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)"
        r"\s+\d{1,2},?\s+\d{4}\b|\b\d{4}\b",
        re.IGNORECASE,
    )

    author_pattern = re.compile(r"^(by\s+|author[s]?:)", re.IGNORECASE)
    footnote_pattern = re.compile(r"^(footnote|endnote|references|bibliography|works cited)\b", re.IGNORECASE)

    for i, line in enumerate(lines):
        if remove_titles and i < 5 and len(line.split()) <= 18:
            continue

        if remove_authors and author_pattern.search(line):
            continue

        if remove_dates and i < 10 and date_pattern.search(line) and len(line.split()) <= 12:
            continue

        if remove_footnotes and footnote_pattern.search(line):
            break

        cleaned.append(line)

    return "\n\n".join(cleaned)


def normalize_spacing(text: str) -> str:
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"\s+([,.!?;:])", r"\1", text)
    return text.strip()


def clean_for_listening(
    text: str,
    remove_titles: bool = False,
    remove_authors: bool = True,
    remove_dates: bool = True,
    remove_citations: bool = True,
    remove_captions_option: bool = True,
    remove_urls_option: bool = True,
    remove_footnotes_option: bool = True,
) -> str:
    if remove_urls_option:
        text = remove_urls(text)

    if remove_citations:
        text = remove_inline_citations(text)

    if remove_footnotes_option:
        text = remove_footnote_markers(text)

    if remove_captions_option:
        text = remove_captions(text)

    text = remove_metadata_lines(
        text,
        remove_titles=remove_titles,
        remove_authors=remove_authors,
        remove_dates=remove_dates,
        remove_footnotes=remove_footnotes_option,
    )

    return normalize_spacing(text)


# -----------------------------
# Reading and audio functions
# -----------------------------

def start_from_phrase(text: str, phrase: str) -> str:
    if not phrase.strip():
        return text

    index = text.lower().find(phrase.lower().strip())
    if index == -1:
        return text

    return text[index:]


def chunk_text(text: str, max_chars: int = 3500) -> List[str]:
    sentences = re.split(r"(?<=[.!?])\s+", text)
    chunks = []
    current = ""

    for sentence in sentences:
        if len(current) + len(sentence) <= max_chars:
            current += sentence + " "
        else:
            if current.strip():
                chunks.append(current.strip())
            current = sentence + " "

    if current.strip():
        chunks.append(current.strip())

    return chunks


async def generate_edge_audio_async(text: str, output_path: str, voice: str, rate: str) -> None:
    communicate = edge_tts.Communicate(text=text, voice=voice, rate=rate)
    await communicate.save(output_path)


def generate_edge_audio(text: str, voice: str, rate: str) -> str:
    """Generate MP3 audio using Edge TTS and return the file path."""
    if not EDGE_TTS_AVAILABLE:
        raise RuntimeError("edge-tts is not installed. Run: pip install edge-tts")

    if not text.strip():
        raise ValueError("There is no text to convert to audio.")

    output_dir = Path(tempfile.gettempdir()) / "human_reader_audio"
    output_dir.mkdir(exist_ok=True)
    output_path = output_dir / "human_reader_output.mp3"

    safe_text = text.strip()

    # Edge TTS works best with reasonably sized text. For long documents, we combine chunks.
    # If this becomes unstable with very long documents, the next step is generating multiple MP3s.
    asyncio.run(generate_edge_audio_async(safe_text, str(output_path), voice, rate))
    return str(output_path)


def speak_with_pyttsx3(text: str, rate: int = 165, volume: float = 1.0, voice_index: int = 0) -> None:
    """Backup local TTS. This is blocking and does not support reliable Streamlit stop controls."""
    if not PYTTSX3_AVAILABLE:
        raise RuntimeError("pyttsx3 is not installed. Run: pip install pyttsx3")

    engine = pyttsx3.init()
    engine.setProperty("rate", rate)
    engine.setProperty("volume", volume)

    voices = engine.getProperty("voices")
    if voices and 0 <= voice_index < len(voices):
        engine.setProperty("voice", voices[voice_index].id)

    for chunk in chunk_text(text):
        engine.say(chunk)
        engine.runAndWait()

    engine.stop()


def get_available_local_voices() -> List[str]:
    if not PYTTSX3_AVAILABLE:
        return ["0: pyttsx3 not installed"]

    try:
        engine = pyttsx3.init()
        voices = engine.getProperty("voices")
        return [f"{i}: {voice.name}" for i, voice in enumerate(voices)]
    except Exception:
        return ["0: Default system voice"]


# -----------------------------
# Streamlit helpers
# -----------------------------

def save_uploaded_file(uploaded_file) -> Optional[str]:
    if uploaded_file is None:
        return None

    suffix = Path(uploaded_file.name).suffix.lower()
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(uploaded_file.read())
        return tmp.name


def extract_uploaded_file(file_path: str) -> str:
    suffix = Path(file_path).suffix.lower()

    if suffix == ".pdf":
        return extract_text_from_pdf(file_path)
    if suffix == ".docx":
        return extract_text_from_docx(file_path)
    if suffix in [".txt", ".md"]:
        return extract_text_from_txt(file_path)

    raise ValueError("Unsupported file type. Please use PDF, DOCX, TXT, or MD.")


def reset_current_item() -> None:
    keys_to_clear = [
        "raw_text",
        "cleaned_text",
        "audio_path",
        "last_reading_text",
        "start_phrase",
    ]
    for key in keys_to_clear:
        if key in st.session_state:
            del st.session_state[key]


def main():
    st.set_page_config(page_title="Human Reader MVP", layout="wide")

    st.title("Human Reader MVP")
    st.caption("An AAEL Labs prototype for cleaner, more human listening from webpages and documents.")

    st.info(
        "Use the browser audio player for pause, stop, replay, and seek controls. "
        "This avoids the Streamlit run-loop error caused by clicking live pyttsx3 reading more than once."
    )

    st.sidebar.header("Listening Cleanup")
    remove_titles = st.sidebar.checkbox("Skip likely titles", value=False)
    remove_authors = st.sidebar.checkbox("Skip author names", value=True)
    remove_dates = st.sidebar.checkbox("Skip publication dates", value=True)
    remove_citations = st.sidebar.checkbox("Remove inline citations", value=True)
    remove_captions_option = st.sidebar.checkbox("Remove captions and source notes", value=True)
    remove_urls_option = st.sidebar.checkbox("Remove URLs", value=True)
    remove_footnotes_option = st.sidebar.checkbox("Remove footnotes/references", value=True)

    st.sidebar.header("Voice Settings")

    st.sidebar.markdown(
        """
        Recommended voices for the most human-sounding experience:

        - JennyNeural → warm and natural
        - AriaNeural → polished audiobook style
        - AvaNeural → conversational and softer
        - GuyNeural → clean professional male voice
        - SoniaNeural → smoother UK narration
        - NatashaNeural → relaxed Australian narration
        """
    )
    reading_engine = st.sidebar.radio(
        "Reading engine",
        ["Browser audio with natural voice", "Backup local live voice"],
        index=0,
    )

    edge_voice = st.sidebar.selectbox(
        "Natural voice",
        [
            # Premium sounding female voices
            "en-US-JennyNeural",
            "en-US-AriaNeural",
            "en-US-AvaNeural",
            "en-US-EmmaNeural",
            "en-US-MichelleNeural",
            "en-US-AnaNeural",

            # Premium sounding male voices
            "en-US-GuyNeural",
            "en-US-DavisNeural",
            "en-US-AndrewNeural",
            "en-US-BrianNeural",
            "en-US-ChristopherNeural",
            "en-US-EricNeural",

            # Conversational / softer voices
            "en-US-NancyNeural",
            "en-US-SaraNeural",
            "en-US-TonyNeural",
            "en-US-SteffanNeural",

            # UK voices
            "en-GB-SoniaNeural",
            "en-GB-RyanNeural",
            "en-GB-LibbyNeural",
            "en-GB-ThomasNeural",

            # Australian voices
            "en-AU-NatashaNeural",
            "en-AU-WilliamNeural",

            # Canadian voices
            "en-CA-ClaraNeural",
            "en-CA-LiamNeural",
        ],
        index=0,
    )

    edge_rate_label = st.sidebar.selectbox(
        "Natural voice speed",
        ["Slower", "Normal", "Faster"],
        index=1,
    )
    edge_rate_map = {
        "Slower": "-15%",
        "Normal": "+0%",
        "Faster": "+15%",
    }
    edge_rate = edge_rate_map[edge_rate_label]

    local_voices = get_available_local_voices()
    selected_local_voice = st.sidebar.selectbox("Backup local voice", local_voices)
    local_voice_index = int(selected_local_voice.split(":")[0]) if ":" in selected_local_voice else 0
    local_rate = st.sidebar.slider("Backup local voice speed", min_value=100, max_value=240, value=165, step=5)
    local_volume = st.sidebar.slider("Backup local volume", min_value=0.1, max_value=1.0, value=1.0, step=0.1)

    st.subheader("Input")
    source_type = st.radio(
        "Choose input type",
        ["Webpage URL", "Upload document", "Paste text"],
        horizontal=True,
    )

    raw_text = ""

    if source_type == "Webpage URL":
        url = st.text_input("Paste webpage URL")
        if st.button("Extract webpage text") and url:
            try:
                with st.spinner("Extracting webpage text..."):
                    raw_text = extract_text_from_webpage(url)
                st.session_state["raw_text"] = raw_text
                st.session_state.pop("audio_path", None)
            except Exception as e:
                st.error(f"Could not extract webpage text: {e}")

    elif source_type == "Upload document":
        uploaded_file = st.file_uploader("Upload PDF, DOCX, TXT, or MD", type=["pdf", "docx", "txt", "md"])
        if st.button("Extract document text") and uploaded_file:
            try:
                with st.spinner("Extracting document text..."):
                    file_path = save_uploaded_file(uploaded_file)
                    raw_text = extract_uploaded_file(file_path)
                st.session_state["raw_text"] = raw_text
                st.session_state.pop("audio_path", None)
            except Exception as e:
                st.error(f"Could not extract document text: {e}")

    else:
        pasted_text = st.text_area("Paste text here", height=250)
        if st.button("Use pasted text") and pasted_text.strip():
            st.session_state["raw_text"] = pasted_text
            st.session_state.pop("audio_path", None)

    raw_text = st.session_state.get("raw_text", "")

    if raw_text:
        cleaned_text = clean_for_listening(
            raw_text,
            remove_titles=remove_titles,
            remove_authors=remove_authors,
            remove_dates=remove_dates,
            remove_citations=remove_citations,
            remove_captions_option=remove_captions_option,
            remove_urls_option=remove_urls_option,
            remove_footnotes_option=remove_footnotes_option,
        )

        st.session_state["cleaned_text"] = cleaned_text

        st.subheader("Start Reading From")
        start_phrase = st.text_input(
            "Optional: type a phrase where reading should begin",
            key="start_phrase",
        )
        reading_text = start_from_phrase(cleaned_text, start_phrase)

        col1, col2 = st.columns(2)

        with col1:
            st.subheader("Cleaned Listening Text")
            edited_text = st.text_area("You can edit before generating audio", value=reading_text, height=500)

        with col2:
            st.subheader("Original Extracted Text")
            st.text_area("Original", value=raw_text, height=500)

        st.download_button(
            "Download cleaned text",
            data=edited_text,
            file_name="cleaned_listening_text.txt",
            mime="text/plain",
        )

        st.subheader("Listen")

        action_col1, action_col2, action_col3 = st.columns(3)

        with action_col1:
            generate_clicked = st.button("Generate audio / Continue")

        with action_col2:
            new_item_clicked = st.button("Enter new item")

        with action_col3:
            exit_clicked = st.button("Exit / clear session")

        if new_item_clicked or exit_clicked:
            reset_current_item()
            st.rerun()

        if generate_clicked:
            if reading_engine == "Browser audio with natural voice":
                try:
                    with st.spinner("Generating browser audio..."):
                        audio_path = generate_edge_audio(edited_text, voice=edge_voice, rate=edge_rate)
                    st.session_state["audio_path"] = audio_path
                    st.session_state["last_reading_text"] = edited_text
                    st.success("Audio generated. Use the audio player's pause/stop controls to finish early.")
                except Exception as e:
                    st.error(f"Could not generate audio: {e}")
                    st.warning("Try the backup local live voice, or confirm edge-tts is installed with: pip install edge-tts")
            else:
                st.warning(
                    "Backup local live voice is blocking. Do not click the button again while it is reading. "
                    "Use browser audio mode for stop/pause controls."
                )
                try:
                    speak_with_pyttsx3(
                        edited_text,
                        rate=local_rate,
                        volume=local_volume,
                        voice_index=local_voice_index,
                    )
                except RuntimeError as e:
                    st.error(f"The local speech engine is already running: {e}")
                except Exception as e:
                    st.error(f"Could not read aloud: {e}")

        audio_path = st.session_state.get("audio_path")
        if audio_path and Path(audio_path).exists():
            audio_bytes = Path(audio_path).read_bytes()
            st.audio(audio_bytes, format="audio/mp3")
            st.caption(
                "To stop early, pause the audio player or move the progress bar. "
                "Then click Generate audio / Continue, Enter new item, or Exit / clear session."
            )


def running_inside_streamlit() -> bool:
    """Return True when the script is already running inside Streamlit."""
    try:
        from streamlit.runtime.scriptrunner import get_script_run_ctx
        return get_script_run_ctx() is not None
    except Exception:
        return False


def launch_streamlit_from_vscode() -> None:
    """
    Allows the user to press Run in VS Code or run python human_reader_mvp.py directly.
    The script relaunches itself using Streamlit and opens the browser.
    """
    script_path = Path(__file__).resolve()
    url = "http://localhost:8501"

    command = [
        sys.executable,
        "-m",
        "streamlit",
        "run",
        str(script_path),
        "--server.headless=false",
        "--server.port=8501",
    ]

    print("Launching Human Reader MVP in Streamlit...")
    print(f"Opening browser at {url}")
    print("Command:", " ".join(command))

    try:
        webbrowser.open(url)
    except Exception:
        pass

    subprocess.run(command)


if __name__ == "__main__":
    if running_inside_streamlit():
        main()
    else:
        launch_streamlit_from_vscode()
